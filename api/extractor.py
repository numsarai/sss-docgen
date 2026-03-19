"""
extractor.py — PDF text extraction and template field scanning.

extract_pdf_text()     reads every page of a PDF and returns a single string.
scan_template_fields() walks a .docx template and returns every {{field}} name
                       found in paragraphs, tables, headers, and footers.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import IO, List, Union

import pdfplumber
from docx import Document


# ---------------------------------------------------------------------------
# PDF text extraction
# ---------------------------------------------------------------------------

def extract_pdf_text(source: Union[str, Path, IO[bytes]]) -> str:
    """
    Extract all text from *source* (path or file-like object).

    Each page is prefixed with a ``[Page N]`` marker so the LLM can
    reason about document structure.  Empty pages are omitted.
    """
    with pdfplumber.open(source) as pdf:
        pages: List[str] = []
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
            if text.strip():
                pages.append(f"[Page {i}]\n{text.strip()}")

    return "\n\n".join(pages)


# ---------------------------------------------------------------------------
# Template field scanning
# ---------------------------------------------------------------------------

def scan_template_fields(template_path: Union[str, Path]) -> List[str]:
    """
    Return a deduplicated, ordered list of ``{{field}}`` names found in a
    ``.docx`` template.

    Scans:
    * body paragraphs
    * body tables (nested tables included)
    * headers and footers for every section (default, first-page, even-page)
    """
    doc = Document(str(template_path))
    chunks: List[str] = []

    # ── Body paragraphs ────────────────────────────────────────────────────
    for para in doc.paragraphs:
        chunks.append(para.text)

    # ── Body tables ────────────────────────────────────────────────────────
    def _collect_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    chunks.append(para.text)
                for nested in cell.tables:
                    _collect_table(nested)

    for table in doc.tables:
        _collect_table(table)

    # ── Headers + footers ──────────────────────────────────────────────────
    for section in doc.sections:
        for part in (
            section.header,            section.first_page_header, section.even_page_header,
            section.footer,            section.first_page_footer, section.even_page_footer,
        ):
            if part and not part.is_linked_to_previous:
                for para in part.paragraphs:
                    chunks.append(para.text)

    # ── Deduplicate while preserving order ─────────────────────────────────
    seen: List[str] = []
    for m in re.finditer(r"\{\{([^{}]+)\}\}", " ".join(chunks)):
        name = m.group(1)
        if name not in seen:
            seen.append(name)

    return seen
