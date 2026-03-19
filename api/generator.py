"""
generator.py — DOCX + PDF generation pipeline.

generate_docx()   fills a template with variables via docx_engine.
convert_to_pdf()  converts the DOCX to PDF using LibreOffice headless.

LibreOffice is preferred because it handles complex .docx faithfully
(fonts, RTL text, Thai characters).  The binary is located by trying
several common install paths before raising a clear error.
"""

import shutil
import subprocess
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Dict

_BKK = timezone(timedelta(hours=7))

from docx import Document
from docx_engine import DocxEngine

# ---------------------------------------------------------------------------
# Thai numeral conversion
# ---------------------------------------------------------------------------

_THAI_DIGITS = str.maketrans('0123456789', '๐๑๒๓๔๕๖๗๘๙')


def _to_thai(text: str) -> str:
    return text.translate(_THAI_DIGITS)


def _apply_thai_numerals(doc: Document) -> None:
    """Replace every Arabic digit with its Thai equivalent in all text runs."""

    def process_para(para):
        for run in para.runs:
            if any(c.isdigit() for c in run.text):
                run.text = _to_thai(run.text)

    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)
                for nested in cell.tables:
                    process_table(nested)

    for para in doc.paragraphs:
        process_para(para)
    for table in doc.tables:
        process_table(table)

    for section in doc.sections:
        for part in (
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer,
        ):
            if part is not None and not part.is_linked_to_previous:
                for para in part.paragraphs:
                    process_para(para)
                for table in part.tables:
                    process_table(table)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

_PROJECT_ROOT  = Path(__file__).parent.parent
GENERATED_DIR  = _PROJECT_ROOT / "generated"

# Candidate LibreOffice executable names / paths
_SOFFICE_BINS = [
    "soffice",
    "libreoffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/lib/libreoffice/program/soffice",
    "/usr/bin/libreoffice",
    "/usr/bin/soffice",
    "/snap/bin/libreoffice",
]


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _find_soffice() -> str:
    """Return the first usable LibreOffice binary path, or raise."""
    for candidate in _SOFFICE_BINS:
        found = shutil.which(candidate)
        if found:
            return found
        p = Path(candidate)
        if p.exists() and p.is_file():
            return str(p)
    raise RuntimeError(
        "LibreOffice not found.\n"
        "Install from https://www.libreoffice.org/download/ "
        "and ensure `soffice` is on PATH."
    )


def _file_stem(case_id: int) -> str:
    """Unique filename stem: case_<id>_<timestamp> in Bangkok time."""
    ts = datetime.now(_BKK).strftime("%Y%m%d_%H%M%S")
    return f"case_{case_id}_{ts}"


def _resolve_template(template: str) -> Path:
    """
    Accept either an absolute path or a name relative to the project root.
    Raises FileNotFoundError with a clear message if not found.
    """
    p = Path(template)
    if p.is_absolute():
        if not p.exists():
            raise FileNotFoundError(f"Template not found: {p}")
        return p

    # Try relative to project root
    candidate = _PROJECT_ROOT / template
    if candidate.exists():
        return candidate

    raise FileNotFoundError(
        f"Template '{template}' not found.\n"
        f"Tried: {p}, {candidate}"
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_docx(template: str, variables: Dict[str, str], case_id: int) -> Path:
    """
    Fill *template* with *variables* and write the result to
    ``generated/case_<id>_<ts>.docx``.

    Returns the absolute Path of the created file.
    """
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)

    tpl_path = _resolve_template(template)
    out_path = GENERATED_DIR / f"{_file_stem(case_id)}.docx"

    engine = DocxEngine(tpl_path)
    engine.fill(variables)
    _apply_thai_numerals(engine._doc)
    engine.save(out_path)
    return out_path


def convert_to_pdf(docx_path: Path) -> Path:
    """
    Convert *docx_path* to PDF using LibreOffice headless.

    The PDF is written alongside the DOCX with the same stem.
    Returns the absolute Path of the created PDF.

    Raises RuntimeError if LibreOffice is unavailable or conversion fails.
    """
    soffice = _find_soffice()

    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(docx_path.parent),
            str(docx_path),
        ],
        capture_output=True,
        timeout=60,
    )

    if result.returncode != 0:
        stderr = result.stderr.decode(errors="replace")
        stdout = result.stdout.decode(errors="replace")
        raise RuntimeError(
            f"LibreOffice exited with code {result.returncode}.\n"
            f"stderr: {stderr}\nstdout: {stdout}"
        )

    pdf_path = docx_path.with_suffix(".pdf")
    if not pdf_path.exists():
        raise RuntimeError(
            f"Conversion appeared to succeed but PDF not found: {pdf_path}"
        )

    return pdf_path
