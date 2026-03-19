"""
replacer.py — safe {{placeholder}} replacement inside python-docx objects.

Key design
----------
Word stores paragraph text in a sequence of *runs*, each with its own
character formatting (bold, font, colour, …).  A placeholder typed by a
human may be split across several runs because Word inserts run boundaries
invisibly (autocorrect, language detection, IME, style changes, etc.).

Strategy
~~~~~~~~
1. Reconstruct the paragraph's full text from all runs and build a
   *character map*: for every character position → (run_index, pos_in_run).
2. Locate every ``{{…}}`` placeholder with a regex.
3. Process matches **right-to-left** so that editing later runs does not
   shift the character positions used by earlier matches.
4. For a placeholder inside a single run: plain in-place substitution.
5. For a placeholder spanning multiple runs:
   - Write  (prefix + value + suffix-of-last-run)  into the *first* run
     (keeping its formatting).
   - Blank every subsequent run that was part of the placeholder so that
     no stray characters remain.
"""

from __future__ import annotations

import re
from typing import Dict, List, Tuple

from docx.text.paragraph import Paragraph
from docx.table import Table

# Matches  {{anything}}  — no nested braces allowed.
_PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _char_map(paragraph: Paragraph) -> List[Tuple[int, int]]:
    """Return a list of ``(run_index, pos_in_run)`` for every character."""
    mapping: List[Tuple[int, int]] = []
    for run_idx, run in enumerate(paragraph.runs):
        for char_idx in range(len(run.text)):
            mapping.append((run_idx, char_idx))
    return mapping


# ---------------------------------------------------------------------------
# Public helpers used by engine.py
# ---------------------------------------------------------------------------

def replace_in_paragraph(paragraph: Paragraph, replacements: Dict[str, str]) -> None:
    """
    Replace every matching ``{{placeholder}}`` in *paragraph* in-place.

    Formatting of the run that *starts* a placeholder is preserved.
    If a placeholder spans multiple runs the intermediate runs are cleared.
    """
    full_text = "".join(run.text for run in paragraph.runs)

    # Fast bail-out — nothing to do
    if "{{" not in full_text:
        return

    matches = [m for m in _PLACEHOLDER_RE.finditer(full_text) if m.group() in replacements]
    if not matches:
        return

    cmap = _char_map(paragraph)
    runs = paragraph.runs

    # Right-to-left so earlier char positions stay valid after each edit
    for match in reversed(matches):
        value = replacements[match.group()]
        start = match.start()
        end   = match.end() - 1   # inclusive last character index

        if start >= len(cmap) or end >= len(cmap):
            continue

        first_run_idx, pos_in_first = cmap[start]
        last_run_idx,  pos_in_last  = cmap[end]
        pos_in_last += 1  # make exclusive

        if first_run_idx == last_run_idx:
            # ── Simple case: placeholder is entirely within one run ──────
            r = runs[first_run_idx]
            r.text = r.text[:pos_in_first] + value + r.text[pos_in_last:]

        else:
            # ── Cross-run case ───────────────────────────────────────────
            # Anything after the placeholder in the last run must be kept.
            tail = runs[last_run_idx].text[pos_in_last:]

            # Rewrite the first run: prefix + replacement + tail of last run.
            first_run = runs[first_run_idx]
            first_run.text = first_run.text[:pos_in_first] + value + tail

            # Erase every subsequent run that carried placeholder characters.
            for idx in range(first_run_idx + 1, last_run_idx + 1):
                runs[idx].text = ""


def replace_in_table(table: Table, replacements: Dict[str, str]) -> None:
    """Recursively replace placeholders in every cell of *table*."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            for nested in cell.tables:          # nested tables
                replace_in_table(nested, replacements)


def replace_in_container(container, replacements: Dict[str, str]) -> None:
    """Replace placeholders in all paragraphs **and** tables of *container*.

    Works for the document body, headers, and footers (all expose
    ``.paragraphs`` and ``.tables``).
    """
    for paragraph in container.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in container.tables:
        replace_in_table(table, replacements)
